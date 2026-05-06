import os
import re
import io
import requests
from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from werkzeug.utils import secure_filename

try:
    from anthropic import Anthropic
    ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY', '')
    client = Anthropic(api_key=ANTHROPIC_API_KEY) if ANTHROPIC_API_KEY else None
except ImportError:
    client = None

app = Flask(__name__)
CORS(app)

app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['PROCESSED_FOLDER'] = 'processed'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max
ALLOWED_EXTENSIONS = {'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def apply_abnt_formatting(doc: Document) -> Document:
    """Aplica formatação ABNT NBR 14724:2011 ao documento"""

    # Configurar margens (superior/esquerda 3cm, inferior/direita 2cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(3)
        section.left_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.right_margin = Cm(2)

    # Definir estilos
    styles = doc.styles

    # Estilo para corpo do texto
    if 'CorpoABNT' not in styles:
        corpo_style = styles.add_style('CorpoABNT', WD_STYLE_TYPE.PARAGRAPH)
        corpo_font = corpo_style.font
        corpo_font.name = 'Arial'
        corpo_font.size = Pt(12)
        corpo_format = corpo_style.paragraph_format
        corpo_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        corpo_format.space_after = Pt(0)
        corpo_format.space_before = Pt(0)
        corpo_format.first_line_indent = Cm(1.25)  # Recuo de primeira linha

    # Estilo para títulos
    if 'TituloABNT' not in styles:
        titulo_style = styles.add_style('TituloABNT', WD_STYLE_TYPE.PARAGRAPH)
        titulo_font = titulo_style.font
        titulo_font.name = 'Arial'
        titulo_font.size = Pt(12)
        titulo_font.bold = True
        titulo_format = titulo_style.paragraph_format
        titulo_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # Estilo para citações longas (>3 linhas)
    if 'CitacaoLonga' not in styles:
        citacao_style = styles.add_style('CitacaoLonga', WD_STYLE_TYPE.PARAGRAPH)
        citacao_font = citacao_style.font
        citacao_font.name = 'Arial'
        citacao_font.size = Pt(10)
        citacao_format = citacao_style.paragraph_format
        citacao_format.left_indent = Cm(4)  # Recuo de 4cm
        citacao_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Estilo para referências
    if 'Referencia' not in styles:
        ref_style = styles.add_style('Referencia', WD_STYLE_TYPE.PARAGRAPH)
        ref_font = ref_style.font
        ref_font.name = 'Arial'
        ref_font.size = Pt(12)
        ref_format = ref_style.paragraph_format
        ref_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        ref_format.space_after = Pt(6)

    # Aplicar formatação aos parágrafos existentes
    for para in doc.paragraphs:
        text = para.text.strip()

        # Detectar tipo de conteúdo e aplicar estilo apropriado
        if len(text) < 100 and not text.endswith('.'):
            # Provavelmente é título
            para.style = 'TituloABNT' if 'TituloABNT' in styles else 'Heading 1'
        elif text.startswith('REFER') or text.startswith('BIBLIO'):
            para.style = 'TituloABNT' if 'TituloABNT' in styles else 'Heading 1'
        elif para.text.startswith(' ') and len(para.text) > 50:
            # Citação longa
            if 'CitacaoLonga' in styles:
                para.style = 'CitacaoLonga'
        else:
            # Corpo do texto
            if 'CorpoABNT' in styles:
                para.style = 'CorpoABNT'
            else:
                # Aplicar formatação manual
                for run in para.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                para.paragraph_format.first_line_indent = Cm(1.25)

    # Formatação de referências
    in_references = False
    for para in doc.paragraphs:
        text = para.text.upper().strip()
        if 'REFER' in text or 'BIBLIO' in text:
            in_references = True
            continue
        if in_references:
            if 'Referencia' in styles:
                para.style = 'Referencia'
            else:
                for run in para.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                para.paragraph_format.space_after = Pt(6)

    return doc

def validate_article(doi=None, title=None, author=None):
    """Valida se um artigo existe usando APIs externas"""
    results = []

    # Tentar Crossref API
    if doi:
        try:
            crossref_url = f"https://api.crossref.org/works/{doi}"
            response = requests.get(crossref_url, timeout=10)
            if response.status_code == 200:
                data = response.json()
                if data['status'] == 'ok':
                    work = data['message']
                    results.append({
                        'source': 'Crossref',
                        'valid': True,
                        'doi': doi,
                        'title': work.get('title', [''])[0],
                        'author': work.get('author', [{}])[0].get('family', ''),
                        'published': work.get('published-print', {}).get('date-parts', [['']])[0][0],
                        'journal': work.get('container-title', [''])[0],
                        'url': f"https://doi.org/{doi}"
                    })
        except Exception as e:
            pass

    # Tentar Semantic Scholar API
    if title:
        try:
            semscholar_url = "https://api.semanticscholar.org/graph/v1/paper/search"
            params = {'query': title[:100], 'limit': 5}
            response = requests.get(semscholar_url, params=params, timeout=10)
            if response.status_code == 200:
                data = response.json()
                for paper in data.get('data', [])[:3]:
                    results.append({
                        'source': 'Semantic Scholar',
                        'valid': True,
                        'doi': paper.get('externalIds', {}).get('DOI', ''),
                        'title': paper.get('title', ''),
                        'author': ', '.join([a['name'] for a in paper.get('authors', [])[:3]]),
                        'published': str(paper.get('publicationDate', ''))[:4],
                        'journal': paper.get('venue', ''),
                        'url': paper.get('url', ''),
                        'citations': paper.get('citationCount', 0)
                    })
        except Exception as e:
            pass

    # Tentar OpenAlex API (backup)
    if title and len(results) == 0:
        try:
            openalex_url = f"https://api.openalex.org/works?filter=title.search:{title[:50]}"
            response = requests.get(openalex_url, timeout=10)
            if response.status_code == 200:
                data = response.json()
                for work in data.get('results', [])[:3]:
                    results.append({
                        'source': 'OpenAlex',
                        'valid': True,
                        'doi': work.get('doi', ''),
                        'title': work.get('title', ''),
                        'author': ', '.join([a['author']['display_name'] for a in work.get('authorships', [])[:3]]),
                        'published': str(work.get('publication_date', ''))[:4],
                        'journal': work.get('host_venue', {}).get('display_name', ''),
                        'url': work.get('doi', ''),
                        'citations': work.get('cited_by_count', 0)
                    })
        except Exception as e:
            pass

    return results

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'Nenhum arquivo enviado'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Nenhum arquivo selecionado'}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        try:
            # Processar documento
            doc = Document(filepath)
            doc = apply_abnt_formatting(doc)

            # Salvar documento formatado
            output_filename = f"ABNT_{filename}"
            output_path = os.path.join(app.config['PROCESSED_FOLDER'], output_filename)
            doc.save(output_path)

            # Extrair texto para análise
            full_text = '\n'.join([p.text for p in doc.paragraphs])

            return jsonify({
                'success': True,
                'message': 'Documento formatado com sucesso!',
                'download_url': f'/download/{output_filename}',
                'original_filename': filename,
                'preview_text': full_text[:2000]
            })
        except Exception as e:
            return jsonify({'error': f'Erro ao processar documento: {str(e)}'}), 500

    return jsonify({'error': 'Formato de arquivo não suportado'}), 400

@app.route('/download/<filename>')
def download_file(filename):
    return send_file(
        os.path.join(app.config['PROCESSED_FOLDER'], filename),
        as_attachment=True
    )

@app.route('/validate', methods=['POST'])
def validate_article_endpoint():
    data = request.get_json()
    doi = data.get('doi', '')
    title = data.get('title', '')
    author = data.get('author', '')

    if not doi and not title:
        return jsonify({'error': 'Informe DOI ou título do artigo'}), 400

    results = validate_article(doi=doi, title=title, author=author)

    if results:
        return jsonify({
            'found': True,
            'articles': results,
            'message': f'{len(results)} artigo(s) encontrado(s)'
        })
    else:
        return jsonify({
            'found': False,
            'articles': [],
            'message': 'Nenhum artigo encontrado. Verifique os dados informados.'
        })

@app.route('/extract-info', methods=['POST'])
def extract_article_info():
    """Extrai informações de artigo de um documento .docx"""
    if 'file' not in request.files:
        return jsonify({'error': 'Nenhum arquivo enviado'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Nenhum arquivo selecionado'}), 400

    if file and allowed_file(file.filename):
        try:
            doc = Document(file)
            full_text = '\n'.join([p.text for p in doc.paragraphs])

            # Extrair DOI (padrão)
            doi_match = re.search(r'10\.\d{4,9}/[-._;()/:A-Z0-9]+', full_text, re.IGNORECASE)
            doi = doi_match.group(0) if doi_match else ''

            # Tentar extrair título (primeira linha não vazia)
            title = ''
            for para in doc.paragraphs:
                if para.text.strip() and len(para.text.strip()) > 10:
                    title = para.text.strip()
                    break

            # Tentar extrair autor (geralmente após título)
            author = ''
            found_title = False
            for para in doc.paragraphs:
                if found_title and para.text.strip():
                    author = para.text.strip()
                    break
                if para.text.strip() == title:
                    found_title = True

            return jsonify({
                'success': True,
                'doi': doi,
                'title': title[:200] if title else '',
                'author': author[:100] if author else '',
                'preview': full_text[:1000]
            })
        except Exception as e:
            return jsonify({'error': f'Erro ao extrair informações: {str(e)}'}), 500

    return jsonify({'error': 'Formato não suportado'}), 400

@app.route('/chat', methods=['POST'])
def chat():
    """Endpoint para chat com IA assistente"""
    data = request.get_json()
    user_message = data.get('message', '')
    conversation_history = data.get('history', [])

    if not user_message:
        return jsonify({'error': 'Mensagem vazia'}), 400

    # System prompt com contexto sobre o site e ABNT
    system_prompt = """Você é um assistente virtual especializado em ajudar usuários com:
1. Formatação de documentos nas normas ABNT (NBR 14724, NBR 6023, NBR 10520)
2. Validação de artigos científicos
3. Uso deste site (ABNT Formatter)

Seu tom deve ser:
- Amigável e prestativo
- Claro e direto
- Especializado em normas acadêmicas

Informações sobre o site:
- Formata documentos .docx automaticamente nas normas ABNT
- Valida artigos via DOI ou título usando Crossref, Semantic Scholar e OpenAlex
- Margens: 3cm (superior/esquerda), 2cm (inferior/direita)
- Fonte: Arial 12, espaçamento 1.5, recuo 1.25cm

Se não souber algo, seja honesto e sugira consultar as normas oficiais."""

    # Construir mensagens para a API
    messages = []
    for msg in conversation_history[-10:]:  # Últimas 10 mensagens
        messages.append({
            "role": msg['role'],
            "content": msg['content']
        })

    try:
        if client:
            response = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=500,
                system=system_prompt,
                messages=[
                    {"role": "user", "content": user_message}
                ] + messages
            )
            ai_response = response.content[0].text
        else:
            # Fallback quando API key não está configurada
            ai_response = gerar_resposta_fallback(user_message)

        return jsonify({
            'success': True,
            'response': ai_response
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'response': gerar_resposta_fallback(user_message)
        })

def gerar_resposta_fallback(mensagem):
    """Gera respostas básicas quando a API do Claude não está disponível"""
    mensagem = mensagem.lower()

    if 'abnt' in mensagem and 'format' in mensagem:
        return "Para formatar seu documento nas normas ABNT:\n\n1. Arraste seu arquivo .docx na área de upload\n2. Clique em 'Formatar Documento'\n3. Aguarde o processamento\n4. Baixe o documento formatado\n\nO sistema aplicará automaticamente: margens 3cm/2cm, fonte Arial 12, espaçamento 1.5 e recuo de 1.25cm."

    if 'doi' in mensagem or 'validar' in mensagem:
        return "Para validar um artigo:\n\n1. Vá na seção 'Validador de Artigos'\n2. Cole o DOI (ex: 10.1000/xyz123) OU digite o título\n3. Clique em 'Validar Artigo'\n\nO sistema consultará Crossref, Semantic Scholar e OpenAlex."

    if 'margem' in mensagem or 'margens' in mensagem:
        return "As margens nas normas ABNT são:\n• Superior: 3cm\n• Esquerda: 3cm\n• Inferior: 2cm\n• Direita: 2cm"

    if 'fonte' in mensagem:
        return "A fonte recomendada pela ABNT é Arial, tamanho 12 para o corpo do texto. Citações longas (>3 linhas) usam tamanho 10."

    if 'citaç' in mensagem or 'citacao' in mensagem:
        return "Citações com mais de 3 linhas devem ter:\n• Recuo de 4cm da margem esquerda\n• Fonte tamanho 10\n• Espaçamento simples\n• Sem aspas"

    if 'refer' in mensagem:
        return "As referências devem seguir a NBR 6023:\n• Alinhamento à esquerda\n• Espaçamento simples\n• Espaço de 6pt entre referências\n• Ordem alfabética"

    if 'olá' in mensagem or 'oi' in mensagem or 'ola' in mensagem:
        return "Olá! Sou o assistente virtual do ABNT Formatter. Como posso ajudar você hoje? Posso tirar dúvidas sobre formatação ABNT, validação de artigos ou como usar o site."

    if 'obrigad' in mensagem:
        return "Por nada! Estou aqui para ajudar. Se tiver mais dúvidas, é só perguntar!"

    return "Desculpe, não entendi completamente. Posso ajudar com:\n• Formatação ABNT de documentos\n• Validação de artigos por DOI ou título\n• Dúvidas sobre normas ABNT (margens, fonte, citações, referências)\n\nComo posso ajudar?"

if __name__ == '__main__':
    os.makedirs('uploads', exist_ok=True)
    os.makedirs('processed', exist_ok=True)
    app.run(debug=True, host='0.0.0.0', port=5000)
